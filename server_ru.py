import logging
import os
import tempfile
from typing import List

from flask import Flask, request, jsonify, send_file
import numpy as np
import uuid
import asyncio

from wav_io.wav_io import transform_to_wavpcm, load_sound
from wav_io.wav_io import TARGET_SAMPLING_FREQUENCY
from asr.asr import initialize_model_for_speech_recognition
from asr.asr import initialize_model_for_speech_segmentation
from asr.asr import async_transcribe as transcribe_speech
from utils.utils import time_to_str

from docx import Document

speech_to_srt_logger = logging.getLogger(__name__)
logging.basicConfig(format='%(asctime)s : %(levelname)s : %(message)s',
                    level=logging.INFO)
app = Flask(__name__)

RESULTS_FOLDER = "results"
os.makedirs(RESULTS_FOLDER, exist_ok=True)

task_status = {}

FRAME_SIZE = 20
LANGUAGE_NAME = 'ru'

try:
    segmenter = initialize_model_for_speech_segmentation(LANGUAGE_NAME)
except BaseException as ex:
    err_msg = str(ex)
    speech_to_srt_logger.error(err_msg)
    raise
speech_to_srt_logger.info('The Wav2Vec2-based segmenter is loaded.')

try:
    asr = initialize_model_for_speech_recognition(LANGUAGE_NAME)
except BaseException as ex:
    err_msg = str(ex)
    speech_to_srt_logger.error(err_msg)
    raise
speech_to_srt_logger.info('The Whisper-based ASR is initialized.')


@app.route('/ready')
def ready():
    return 'OK'


@app.route('/transcribe', methods=['POST'])
async def transcribe():
    task_id = str(uuid.uuid4())
    task_status[task_id] = {"status": "Not ready"}
    speech_to_srt_logger.info('~~~Recognition process started~~~')
    if 'audio' not in request.files:
        speech_to_srt_logger.error('400: No audiofile part in the request')
        task_status[task_id] = jsonify(
            {"status": "Error", "status_code": 400, "message": "No audiofile part in the request"})
        return task_status[task_id]
    file = request.files['audio']
    if file.filename == '':
        speech_to_srt_logger.error('400: No audio file provided for upload')
        task_status[task_id] = jsonify(
            {"status": "Error", "status_code": 400, "message": "No audio file provided for upload"})
        return task_status[task_id]

    point_pos = file.filename.rfind('.')
    if point_pos > 0:
        src_file_ext = file.filename[(point_pos + 1):]
    else:
        src_file_ext = ''
    if len(src_file_ext) == 0:
        speech_to_srt_logger.error('400: Unknown type of the file provided for upload')
        task_status[task_id] = jsonify(
            {"status": "Error", "status_code": 400, "message": "Unknown type of the file provided for upload"})
        return task_status[task_id]
    tmp_audio_name = ''
    tmp_wav_name = ''
    err_msg = ''
    input_sound = None
    try:
        with tempfile.NamedTemporaryFile(mode='wb', delete=False, suffix='.' + src_file_ext) as fp:
            tmp_audio_name = fp.name
        file.save(tmp_audio_name)
        speech_to_srt_logger.info(f'The sound "{file.filename}" is saved to the "{tmp_audio_name}".')
        with tempfile.NamedTemporaryFile(mode='wb', delete=False, suffix='.wav') as fp:
            tmp_wav_name = fp.name
        try:
            transform_to_wavpcm(tmp_audio_name, tmp_wav_name)
        except BaseException as ex:
            err_msg = str(ex)
            speech_to_srt_logger.error(err_msg)
        if len(err_msg) == 0:
            speech_to_srt_logger.info(f'The sound "{file.filename}" is converted to the "{tmp_wav_name}".')
            try:
                input_sound = load_sound(tmp_wav_name)
            except BaseException as ex:
                err_msg = str(ex)
                speech_to_srt_logger.error(err_msg)
            if len(err_msg) == 0:
                speech_to_srt_logger.info(f'The sound is "{tmp_wav_name}" is loaded.')
    finally:
        if os.path.isfile(tmp_audio_name):
            os.remove(tmp_audio_name)
            speech_to_srt_logger.info(f'The sound is "{tmp_audio_name}" is removed.')
        if os.path.isfile(tmp_wav_name):
            os.remove(tmp_wav_name)
            speech_to_srt_logger.info(f'The sound is "{tmp_wav_name}" is removed.')
    if (len(err_msg) > 0) or (input_sound is None):
        if len(err_msg) == 0:
            err_msg = 'Audio file is empty.'
            speech_to_srt_logger.error('400: ' + err_msg)
            task_status[task_id]["message"] = err_msg
        else:
            task_status[task_id]["message"] = err_msg
        task_status[task_id]["status"] = "Error"
        task_status[task_id]["status_code"] = 400
        return jsonify(task_status[task_id])

    if not isinstance(input_sound, np.ndarray):
        speech_to_srt_logger.info(f'The sound "{file.filename}" is stereo.')
        input_sound = (input_sound[0] + input_sound[1]) / 2.0
    speech_to_srt_logger.info(f'The total duration of the new sound "{file.filename}" is '
                              f'{time_to_str(input_sound.shape[0] / TARGET_SAMPLING_FREQUENCY)}.')

    if input_sound is None:
        speech_to_srt_logger.info(f'The sound "{file.filename}" is empty.')
    else:
        speech_to_srt_logger.error(task_id)
        asyncio.create_task(create_result_file(input_sound, segmenter, asr, FRAME_SIZE, task_id))

    return jsonify({"task_id": task_id})


async def create_result_file(input_sound, segmenter, asr, FRAME_SIZE, task_id):
    texts_with_timestamps = await transcribe_speech(input_sound, segmenter, asr, FRAME_SIZE)
    output_filename = task_id + ".docx"
    doc = Document()
    for start_time, end_time, sentence_text in texts_with_timestamps:
        line = f"{start_time:.2f} - {end_time:.2f} - {sentence_text}"
        doc.add_paragraph(line)
        doc.add_paragraph("")

    result_path = os.path.join(RESULTS_FOLDER, output_filename)
    doc.save(result_path)

    task_status[task_id] = jsonify({"status": "Ready", "status_code": 200, "result_path": result_path})


@app.route("/status/<task_id>", methods=["GET"])
def get_status(task_id):
    status = task_status.get(task_id, None)
    if status is None:
        return jsonify({"error": "Task not found"}), 404
    return status


@app.route("/download_result/<task_id>", methods=["GET"])
def download_result(task_id):
    result_file = f"results/{task_id}.docx"
    if not os.path.isfile(result_file):
        return jsonify({"error": "Result not ready or task not found"}), 404

    return send_file(result_file, as_attachment=True)


if __name__ == '__main__':
    app.run(host='0.0.0.0', port=8040)
