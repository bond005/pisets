from argparse import ArgumentParser
import codecs
import logging
import os
import sys
import tempfile

from docx import Document
import numpy as np

from wav_io.wav_io import transform_to_wavpcm, load_sound
from wav_io.wav_io import TARGET_SAMPLING_FREQUENCY
from asr.asr import initialize_model_for_speech_recognition
from asr.asr import initialize_model_for_speech_classification
from asr.asr import initialize_model_for_speech_segmentation
from asr.asr import transcribe, check_language
from asr.asr import asr_logger
from utils.utils import time_to_str


speech_to_srt_logger = logging.getLogger(__name__)


def main():
    parser = ArgumentParser()
    parser.add_argument('--lang', dest='language', type=str, required=False, default='ru',
                        help='The language of input speech (Russian or English).')
    parser.add_argument('-i', '--input', dest='input_name', type=str, required=True,
                        help='The input sound file name.')
    parser.add_argument('-m', '--model', dest='model_dir', type=str, required=False, default=None,
                        help='The path to directory with Wav2Vec2, AudioTransformer and Whisper.')
    parser.add_argument('-o', '--output', dest='output_name', type=str, required=True,
                        help='The output DocX file name.')
    args = parser.parse_args()

    language_name = check_language(args.language)

    if args.model_dir is None:
        wav2vec2_path = None
        audiotransformer_path = None
        whisper_path = None
    else:
        model_dir = os.path.normpath(args.model_dir)
        if not os.path.isdir(model_dir):
            err_msg = f'The directory "{model_dir}" does not exist!'
            speech_to_srt_logger.error(err_msg)
            raise IOError(err_msg)
        wav2vec2_path = os.path.join(model_dir, 'wav2vec2')
        if not os.path.isdir(wav2vec2_path):
            err_msg = f'The directory "{wav2vec2_path}" does not exist!'
            speech_to_srt_logger.error(err_msg)
            raise IOError(err_msg)
        audiotransformer_path = os.path.join(model_dir, 'ast')
        if not os.path.isdir(audiotransformer_path):
            err_msg = f'The directory "{audiotransformer_path}" does not exist!'
            speech_to_srt_logger.error(err_msg)
            raise IOError(err_msg)
        whisper_path = os.path.join(model_dir, 'whisper')
        if not os.path.isdir(whisper_path):
            err_msg = f'The directory "{whisper_path}" does not exist!'
            speech_to_srt_logger.error(err_msg)
            raise IOError(err_msg)

    audio_fname = os.path.normpath(args.input_name)
    if not os.path.isfile(audio_fname):
        err_msg = f'The file "{audio_fname}" does not exist!'
        speech_to_srt_logger.error(err_msg)
        raise IOError(err_msg)

    output_docx_fname = os.path.normpath(args.output_name)
    output_docx_dir = os.path.dirname(output_docx_fname)
    if len(output_docx_dir) > 0:
        if not os.path.isdir(output_docx_dir):
            err_msg = f'The directory "{output_docx_dir}" does not exist!'
            speech_to_srt_logger.error(err_msg)
            raise IOError(err_msg)
    if len(os.path.basename(output_docx_fname).strip()) == 0:
        err_msg = f'The file name "{output_docx_fname}" is incorrect!'
        speech_to_srt_logger.error(err_msg)
        raise IOError(err_msg)

    if os.path.basename(output_docx_fname) == os.path.basename(audio_fname):
        err_msg = f'The input audio and the output DocX file have a same names! ' \
                  f'{os.path.basename(audio_fname)} = {os.path.basename(output_docx_fname)}'
        speech_to_srt_logger.error(err_msg)
        raise IOError(err_msg)

    tmp_wav_name = ''
    try:
        with tempfile.NamedTemporaryFile(mode='wb', delete=False, suffix='.wav') as fp:
            tmp_wav_name = fp.name
        try:
            transform_to_wavpcm(audio_fname, tmp_wav_name)
        except BaseException as ex:
            err_msg = str(ex)
            speech_to_srt_logger.error(err_msg)
            raise
        speech_to_srt_logger.info(f'The sound "{audio_fname}" is converted to the "{tmp_wav_name}".')
        try:
            input_sound = load_sound(tmp_wav_name)
        except BaseException as ex:
            err_msg = str(ex)
            speech_to_srt_logger.error(err_msg)
            raise
        speech_to_srt_logger.info(f'The sound is "{tmp_wav_name}" is loaded.')
    finally:
        if os.path.isfile(tmp_wav_name):
            os.remove(tmp_wav_name)
            speech_to_srt_logger.info(f'The sound is "{tmp_wav_name}" is removed.')

    if input_sound is None:
        speech_to_srt_logger.info(f'The sound "{audio_fname}" is empty.')
        texts_with_timestamps = []
    else:
        if not isinstance(input_sound, np.ndarray):
            speech_to_srt_logger.info(f'The sound "{audio_fname}" is stereo.')
            input_sound = (input_sound[0] + input_sound[1]) / 2.0
        speech_to_srt_logger.info(f'The total duration of the sound "{audio_fname}" is '
                                  f'{time_to_str(input_sound.shape[0] / TARGET_SAMPLING_FREQUENCY)}.')

        try:
            segmenter = initialize_model_for_speech_segmentation(language_name, model_info=wav2vec2_path)
        except BaseException as ex:
            err_msg = str(ex)
            speech_to_srt_logger.error(err_msg)
            raise
        speech_to_srt_logger.info('The Wav2Vec2-based segmenter is loaded.')

        try:
            vad = initialize_model_for_speech_classification(model_info=audiotransformer_path)
        except BaseException as ex:
            err_msg = str(ex)
            speech_to_srt_logger.error(err_msg)
            raise
        speech_to_srt_logger.info('The AST-based voice activity detector is loaded.')

        try:
            asr = initialize_model_for_speech_recognition(language_name, model_info=whisper_path)
        except BaseException as ex:
            err_msg = str(ex)
            speech_to_srt_logger.error(err_msg)
            raise
        speech_to_srt_logger.info('The Whisper-based ASR is initialized.')

        texts_with_timestamps = transcribe(input_sound, segmenter, vad, asr, min_segment_size=1, max_segment_size=20)

    doc = Document()
    for start_time, end_time, sentence_text in texts_with_timestamps:
        line = f'{time_to_str(start_time)} - {time_to_str(end_time)} - {sentence_text}'
        doc.add_paragraph(line)
        doc.add_paragraph('')
    doc.save(output_docx_fname)


if __name__ == '__main__':
    speech_to_srt_logger.setLevel(logging.INFO)
    asr_logger.setLevel(logging.INFO)
    fmt_str = '%(filename)s[LINE:%(lineno)d]# %(levelname)-8s ' \
              '[%(asctime)s]  %(message)s'
    formatter = logging.Formatter(fmt_str)
    stdout_handler = logging.StreamHandler(sys.stdout)
    stdout_handler.setFormatter(formatter)
    speech_to_srt_logger.addHandler(stdout_handler)
    asr_logger.addHandler(stdout_handler)
    file_handler = logging.FileHandler('speech_to_docx.log')
    file_handler.setFormatter(formatter)
    speech_to_srt_logger.addHandler(file_handler)
    asr_logger.addHandler(file_handler)
    main()
